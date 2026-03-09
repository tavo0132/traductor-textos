import os
import sys
from pathlib import Path
from pypdf import PdfReader
from deep_translator import GoogleTranslator
from docx import Document
from docx2pdf import convert
import time

def traducir_pdf_a_word_y_pdf(ruta_pdf, ruta_salida_sin_extension):
    """
    Traduce un PDF de inglés a español, genera .docx y luego .pdf.
    """
    print(f"[PDF] Procesando PDF: {ruta_pdf}")

    # 1. Extraer texto
    try:
        texto_completo = ""
        reader = PdfReader(ruta_pdf)
        print(f"[INFO] Total de paginas: {len(reader.pages)}")
        
        for i, pagina in enumerate(reader.pages):
            texto_pagina = pagina.extract_text()
            if texto_pagina:
                texto_completo += texto_pagina + "\n\n"
            print(f"   [OK] Pagina {i+1}/{len(reader.pages)} extraida")
        
        if not texto_completo.strip():
            print("[ERROR] No se pudo extraer texto del PDF")
            return False
            
    except Exception as e:
        print(f"[ERROR] Extrayendo PDF: {e}")
        return False

    # 2. Traducir por fragmentos
    print("\n[TRADUCCION] Iniciando traduccion...")
    traductor = GoogleTranslator(source='en', target='es')
    parrafos = texto_completo.split('\n')
    fragmentos = []
    frag_actual = ""
    
    for p in parrafos:
        if len(frag_actual) + len(p) < 4500:
            frag_actual += p + "\n"
        else:
            if frag_actual:
                fragmentos.append(frag_actual)
            frag_actual = p + "\n"
    
    if frag_actual:
        fragmentos.append(frag_actual)

    print(f"[INFO] Traduciendo {len(fragmentos)} fragmentos...")
    texto_traducido = ""
    
    for i, frag in enumerate(fragmentos):
        try:
            if frag.strip():
                print(f"   [PROCESANDO] Fragmento {i+1}/{len(fragmentos)}...", end="\r")
                traducido = traductor.translate(frag)
                texto_traducido += traducido + "\n\n"
                print(f"   [OK] Fragmento {i+1}/{len(fragmentos)} completado     ")
            time.sleep(1)  # Pausa para evitar bloqueos
        except Exception as e:
            print(f"   [ERROR] Fragmento {i+1}: {e}")
            texto_traducido += "[ERROR EN TRADUCCION]\n\n"

    # 3. Crear .docx traducido
    try:
        doc = Document()
        doc.add_heading('Documento Traducido al Español', level=1)
        doc.add_paragraph("(Traduccion automatica de: Project 1 (servers plus).pdf)")
        doc.add_paragraph("")  # Espaciador
        
        for parrafo in texto_traducido.split('\n'):
            if parrafo.strip():
                doc.add_paragraph(parrafo)
        
        ruta_docx = ruta_salida_sin_extension + ".docx"
        doc.save(ruta_docx)
        print(f"\n[OK] DOCX TRADUCIDO guardado en: {ruta_docx}")
    except Exception as e:
        print(f"[ERROR] Creando DOCX traducido: {e}")
        return False

    # 4. Convertir DOCX traducido a PDF
    try:
        ruta_pdf_traducido = ruta_salida_sin_extension + ".pdf"
        convert(ruta_docx, ruta_pdf_traducido)
        print(f"[OK] PDF TRADUCIDO generado en: {ruta_pdf_traducido}")
    except Exception as e:
        print(f"[ADVERTENCIA] No se pudo generar PDF automaticamente: {e}")
        print("   Intenta instalar LibreOffice o convierte manualmente\n")

    # 5. Copiar PDF original a Word
    print("\n[INFO] Creando copia del PDF original en formato Word...")
    try:
        doc_original = Document()
        doc_original.add_heading('PDF Original (en Ingles)', level=1)
        doc_original.add_paragraph("Copia del archivo: Project 1 (servers plus).pdf")
        doc_original.add_paragraph("")
        
        # Extraer y agregar texto original
        reader_original = PdfReader(ruta_pdf)
        for pagina in reader_original.pages:
            texto = pagina.extract_text()
            if texto:
                doc_original.add_paragraph(texto)
        
        ruta_docx_original = ruta_salida_sin_extension + "_ORIGINAL.docx"
        doc_original.save(ruta_docx_original)
        print(f"[OK] PDF ORIGINAL (en Word) guardado en: {ruta_docx_original}")
        
        # Intentar convertir a PDF tambien
        try:
            ruta_pdf_original = ruta_salida_sin_extension + "_ORIGINAL.pdf"
            convert(ruta_docx_original, ruta_pdf_original)
            print(f"[OK] PDF ORIGINAL (copia) generado en: {ruta_pdf_original}")
        except:
            print(f"[ADVERTENCIA] No se pudo generar PDF del original, pero .docx esta disponible")
    
    except Exception as e:
        print(f"[ADVERTENCIA] Error creando copia en Word: {e}")

    print("\n" + "="*70)
    print("[EXITO] PROCESO COMPLETADO")
    print("="*70)
    return True

if __name__ == "__main__":
    # Definir rutas
    carpeta_trabajo = Path(__file__).parent
    carpeta_entrada = carpeta_trabajo / "archivo_para_traducir"
    carpeta_salida = carpeta_trabajo / "salida"
    
    pdf_entrada = carpeta_entrada / "Project 1 (servers plus).pdf"
    nombre_base = carpeta_salida / "Proyecto_Traducido"
    
    # Verificar que el PDF existe
    if not pdf_entrada.exists():
        print(f"[ERROR] No se encontro el PDF en: {pdf_entrada}")
        sys.exit(1)
    
    print(f"[INICIO] PROCESO DE TRADUCCION")
    print(f"[CONFIG] Carpeta de trabajo: {carpeta_trabajo}")
    print(f"[CONFIG] Entrada: {pdf_entrada}")
    print(f"[CONFIG] Salida: {carpeta_salida}\n")
    
    # Ejecutar traducción
    exito = traducir_pdf_a_word_y_pdf(str(pdf_entrada), str(nombre_base))
    
    if exito:
        print("\n[SUCCESS] Traduccion completada exitosamente!")
    else:
        print("\n[ERROR] Error durante el proceso de traduccion")
        sys.exit(1)
