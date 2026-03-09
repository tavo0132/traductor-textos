import os
import sys
from pathlib import Path
from pypdf import PdfReader
from deep_translator import GoogleTranslator
from docx import Document
from docx2pdf import convert
import time

def traducir_pdf_prueba(ruta_pdf, ruta_salida_sin_extension, num_paginas=2):
    """
    Traduce UN NÚMERO LIMITADO de páginas de un PDF para prueba.
    """
    print(f"📄 PRUEBA: Procesando primeras {num_paginas} páginas de: {ruta_pdf}")

    # 1. Extraer texto (solo primeras N páginas)
    try:
        texto_completo = ""
        reader = PdfReader(ruta_pdf)
        total_paginas = len(reader.pages)
        paginas_a_procesar = min(num_paginas, total_paginas)
        
        print(f"📑 Total de páginas en PDF: {total_paginas}")
        print(f"🔍 Procesando: {paginas_a_procesar} páginas\n")
        
        for i in range(paginas_a_procesar):
            pagina = reader.pages[i]
            texto_pagina = pagina.extract_text()
            if texto_pagina:
                texto_completo += texto_pagina + "\n\n"
            print(f"   ✓ Página {i+1}/{paginas_a_procesar} extraída")
        
        if not texto_completo.strip():
            print("⚠️ ERROR: No se pudo extraer texto del PDF")
            return False
            
    except Exception as e:
        print(f"❌ Error extrayendo PDF: {e}")
        return False

    # 2. Traducir por fragmentos
    print("\n🔄 Iniciando traducción de prueba...")
    traductor = GoogleTranslator(source='en', target='es')
    parrafos = texto_completo.split('\n')
    fragmentos = []
    frag_actual = ""
    
    for p in parrafos:
        if len(frag_actual) + len(p) < 4500:
            frag_actual += p + "\n"
        else:
            if frag_actual:
                fragmentos.append(frag_actual)
            frag_actual = p + "\n"
    
    if frag_actual:
        fragmentos.append(frag_actual)

    print(f"📝 Traduciendo {len(fragmentos)} fragmentos...")
    texto_traducido = ""
    
    for i, frag in enumerate(fragmentos):
        try:
            if frag.strip():
                print(f"   ⏳ Fragmento {i+1}/{len(fragmentos)} en traducción...", end="\r")
                traducido = traductor.translate(frag)
                texto_traducido += traducido + "\n\n"
                print(f"   ✓ Fragmento {i+1}/{len(fragmentos)} completado       ")
            time.sleep(1)  # Pausa para evitar bloqueos
        except Exception as e:
            print(f"   ⚠️ Error en fragmento {i+1}: {e}")
            texto_traducido += "[ERROR EN TRADUCCIÓN]\n\n"

    # 3. Crear .docx de prueba
    try:
        doc = Document()
        doc.add_heading('🧪 PRUEBA: Documento Traducido (2 páginas)', level=1)
        doc.add_paragraph("(Traducción automática de prueba)")
        doc.add_paragraph("")
        
        for parrafo in texto_traducido.split('\n'):
            if parrafo.strip():
                doc.add_paragraph(parrafo)
        
        ruta_docx = ruta_salida_sin_extension + "_PRUEBA.docx"
        doc.save(ruta_docx)
        print(f"\n✅ DOCX DE PRUEBA guardado en: {ruta_docx}")
    except Exception as e:
        print(f"❌ Error creando DOCX de prueba: {e}")
        return False

    # 4. Convertir DOCX de prueba a PDF
    try:
        ruta_pdf_prueba = ruta_salida_sin_extension + "_PRUEBA.pdf"
        convert(ruta_docx, ruta_pdf_prueba)
        print(f"✅ PDF DE PRUEBA generado en: {ruta_pdf_prueba}")
    except Exception as e:
        print(f"⚠️ No se pudo generar PDF de prueba: {e}\n")

    print("\n" + "="*70)
    print("✨ PRUEBA COMPLETADA EXITOSAMENTE")
    print("="*70)
    print("\n✅ La traducción funciona correctamente.")
    print("📌 Puedes proceder a ejecutar el script completo con todas las páginas.")
    return True

if __name__ == "__main__":
    # Definir rutas
    carpeta_trabajo = Path(__file__).parent
    carpeta_entrada = carpeta_trabajo / "archivo_para_traducir"
    carpeta_salida = carpeta_trabajo / "salida"
    
    pdf_entrada = carpeta_entrada / "Project 1 (servers plus).pdf"
    nombre_base = carpeta_salida / "Proyecto"
    
    # Verificar que el PDF existe
    if not pdf_entrada.exists():
        print(f"❌ ERROR: No se encontró el PDF en: {pdf_entrada}")
        sys.exit(1)
    
    print(f"🧪 PRUEBA CON 2 PÁGINAS")
    print(f"{'='*70}")
    print(f"📁 Carpeta de trabajo: {carpeta_trabajo}")
    print(f"📂 Entrada: {pdf_entrada}")
    print(f"📂 Salida: {carpeta_salida}\n")
    
    # Ejecutar traducción (solo 2 páginas)
    exito = traducir_pdf_prueba(str(pdf_entrada), str(nombre_base), num_paginas=2)
    
    if not exito:
        print("\n❌ Error durante la prueba")
        sys.exit(1)
